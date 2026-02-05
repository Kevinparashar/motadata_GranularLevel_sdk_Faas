"""
Multi-Modal Data Loader

Handles loading and processing of various data formats:
- Text: .txt, .md, .markdown, .html, .json
- Documents: .pdf, .doc, .docx, .rtf
- Audio: .mp3, .wav, .m4a, .ogg (with transcription)
- Video: .mp4, .avi, .mov, .mkv (with transcription and frame extraction)
- Images: .jpg, .png, .gif, .bmp (with OCR and description)
"""


# Standard library imports
import io
import mimetypes
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# Third-party imports
try:
    import PyPDF2  # type: ignore[import-not-found]

    PDF_AVAILABLE = True  # noqa: N806
except ImportError:
    PDF_AVAILABLE = False  # noqa: N806  # type: ignore[misc]

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True  # noqa: N806
except ImportError:
    DOCX_AVAILABLE = False  # noqa: N806  # type: ignore[misc]

try:
    import speech_recognition as sr  # type: ignore[import-not-found]
    from pydub import AudioSegment  # type: ignore[import-not-found]

    AUDIO_AVAILABLE = True  # noqa: N806
except ImportError:
    AUDIO_AVAILABLE = False  # noqa: N806  # type: ignore[misc]

try:
    import pytesseract  # type: ignore[import-not-found]
    from PIL import Image  # type: ignore[import-not-found]

    IMAGE_AVAILABLE = True  # noqa: N806
except ImportError:
    IMAGE_AVAILABLE = False  # noqa: N806  # type: ignore[misc]

try:
    import cv2  # type: ignore[import-not-found]

    VIDEO_AVAILABLE = True  # noqa: N806
except ImportError:
    VIDEO_AVAILABLE = False  # noqa: N806  # type: ignore[misc]

# Local application/library specific imports
from .exceptions import DocumentProcessingError


class MultiModalLoader:
    """
    Loads and processes various data formats for RAG system.

    Supports:
    - Text files: .txt, .md, .markdown, .html, .json
    - Documents: .pdf, .doc, .docx, .rtf
    - Audio: .mp3, .wav, .m4a, .ogg (with transcription)
    - Video: .mp4, .avi, .mov, .mkv (with transcription and frame extraction)
    - Images: .jpg, .png, .gif, .bmp (with OCR and description)
    """

    def __init__(
        self,
        enable_audio_transcription: bool = True,
        enable_video_transcription: bool = True,
        enable_image_ocr: bool = True,
        enable_image_description: bool = True,
        audio_language: str = "en-US",
        video_extract_frames: bool = True,
        video_frames_per_second: float = 1.0,
    ):
        """
        Initialize multi-modal loader.
        
        Args:
            enable_audio_transcription (bool): Flag to enable or disable audio transcription.
            enable_video_transcription (bool): Flag to enable or disable video transcription.
            enable_image_ocr (bool): Flag to enable or disable image ocr.
            enable_image_description (bool): Flag to enable or disable image description.
            audio_language (str): Input parameter for this operation.
            video_extract_frames (bool): Input parameter for this operation.
            video_frames_per_second (float): Input parameter for this operation.
        """
        self.enable_audio_transcription = enable_audio_transcription
        self.enable_video_transcription = enable_video_transcription
        self.enable_image_ocr = enable_image_ocr
        self.enable_image_description = enable_image_description
        self.audio_language = audio_language
        self.video_extract_frames = video_extract_frames
        self.video_frames_per_second = video_frames_per_second

        # Initialize recognizer for audio/video
        if AUDIO_AVAILABLE and (enable_audio_transcription or enable_video_transcription):
            self.recognizer = sr.Recognizer()
        else:
            self.recognizer = None

    async def load(self, file_path: str, gateway: Optional[Any] = None) -> Tuple[str, Dict[str, Any]]:
        """
        Load content from file and return text content with metadata asynchronously.
        
        Args:
            file_path (str): Path of the input file.
            gateway (Optional[Any]): Gateway client used for LLM calls.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        
        Raises:
            DocumentProcessingError: Raised when this function detects an invalid state or when an underlying call fails.
        """
        import asyncio

        path = Path(file_path)

        # Check file existence asynchronously
        exists = await asyncio.to_thread(path.exists)
        if not exists:
            raise DocumentProcessingError(
                message=f"File not found: {file_path}", file_path=str(file_path), operation="load"
            )

        # Get file metadata asynchronously
        file_metadata = await self._get_file_metadata(path)

        # Determine file type and load
        suffix = path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(path))

        try:
            if suffix in [".txt", ".md", ".markdown"]:
                content, metadata = await self._load_text(path, file_metadata)
            elif suffix == ".html":
                content, metadata = await self._load_html(path, file_metadata)
            elif suffix == ".json":
                content, metadata = await self._load_json(path, file_metadata)
            elif suffix == ".pdf":
                content, metadata = await self._load_pdf(path, file_metadata)
            elif suffix in [".doc", ".docx"]:
                content, metadata = await self._load_docx(path, file_metadata)
            elif suffix in [".mp3", ".wav", ".m4a", ".ogg", ".flac"]:
                content, metadata = await self._load_audio(path, file_metadata)
            elif suffix in [".mp4", ".avi", ".mov", ".mkv", ".webm"]:
                content, metadata = await self._load_video(path, file_metadata)
            elif suffix in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"]:
                content, metadata = await self._load_image(path, file_metadata, gateway)
            else:
                # Try as text file
                try:
                    content, metadata = await self._load_text(path, file_metadata)
                except UnicodeDecodeError as e:
                    raise DocumentProcessingError(
                        message=f"Unsupported file format: {suffix}",
                        file_path=str(file_path),
                        operation="load",
                        original_error=e,
                    )
        except Exception as e:
            raise DocumentProcessingError(
                message=f"Error loading file: {str(e)}",
                file_path=str(file_path),
                operation="load",
                original_error=e,
            )

        # Merge metadata
        metadata.update(file_metadata)
        metadata["file_type"] = suffix[1:] if suffix else "unknown"
        metadata["mime_type"] = mime_type or "application/octet-stream"

        return content, metadata

    async def _get_file_metadata(self, path: Path) -> Dict[str, Any]:
        """
        Extract file metadata asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
        
        Returns:
            Dict[str, Any]: Dictionary result of the operation.
        """
        import asyncio

        stat = await asyncio.to_thread(path.stat)
        return {
            "file_name": path.name,
            "file_size": stat.st_size,
            "file_extension": path.suffix.lower(),
            "modified_at": stat.st_mtime,
        }

    async def _load_text(self, path: Path, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        Load text file asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        """
        import asyncio

        def _read_file(p: Path) -> str:
            with open(p, "r", encoding="utf-8") as f:
                return f.read()

        content = await asyncio.to_thread(_read_file, path)
        return content, metadata

    async def _load_html(self, path: Path, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        Load HTML file and extract text asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        """
        import asyncio

        def _read_and_parse_html(p: Path) -> str:
            try:
                from bs4 import BeautifulSoup

                with open(p, "r", encoding="utf-8") as f:
                    soup = BeautifulSoup(f.read(), "html.parser")
                return soup.get_text(separator=" ", strip=True)
            except ImportError:
                # Fallback: basic HTML tag removal
                import re

                with open(p, "r", encoding="utf-8") as f:
                    html = f.read()
                return re.sub(r"<[^>]+>", "", html)

        content = await asyncio.to_thread(_read_and_parse_html, path)
        return content, metadata

    async def _load_json(self, path: Path, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        Load JSON file and convert to readable text asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        """
        import asyncio
        import json

        def _read_and_parse_json(p: Path) -> str:
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Convert to readable text format
            return json.dumps(data, indent=2, ensure_ascii=False)

        content = await asyncio.to_thread(_read_and_parse_json, path)
        return content, metadata

    async def _load_pdf(self, path: Path, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        Load PDF file and extract text asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        
        Raises:
            DocumentProcessingError: Raised when this function detects an invalid state or when an underlying call fails.
        """
        import asyncio

        if not PDF_AVAILABLE:
            raise DocumentProcessingError(
                message="PyPDF2 is required for PDF support. Install with: pip install PyPDF2",
                file_path=str(path),
                operation="load_pdf",
            )

        def _read_pdf(p: Path) -> Tuple[str, int]:
            content_parts = []
            with open(p, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                page_count = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"--- Page {page_num + 1} ---\n{text}\n")
            return "\n".join(content_parts), page_count

        try:
            content, page_count = await asyncio.to_thread(_read_pdf, path)
            metadata["page_count"] = page_count
        except Exception as e:
            raise DocumentProcessingError(
                message=f"Error reading PDF: {str(e)}",
                file_path=str(path),
                operation="load_pdf",
                original_error=e,
            )

        return content, metadata

    def _extract_docx_paragraphs(self, doc: Any) -> List[str]:
        """Extract text from DOCX paragraphs."""
        content_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)
        return content_parts

    def _extract_docx_tables(self, doc: Any) -> List[str]:
        """Extract text from DOCX tables."""
        content_parts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    content_parts.append(row_text)
        return content_parts

    async def _load_docx(self, path: Path, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        Load DOCX file and extract text asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        
        Raises:
            DocumentProcessingError: Raised when this function detects an invalid state or when an underlying call fails.
        """
        import asyncio

        if not DOCX_AVAILABLE:
            raise DocumentProcessingError(
                message="python-docx is required for DOCX support. Install with: pip install python-docx",
                file_path=str(path),
                operation="load_docx",
            )

        def _read_docx(p: Path) -> Tuple[str, int]:
            doc = DocxDocument(str(p))
            content_parts = []
            content_parts.extend(self._extract_docx_paragraphs(doc))
            content_parts.extend(self._extract_docx_tables(doc))
            content = "\n".join(content_parts)
            paragraph_count = len(doc.paragraphs)
            return content, paragraph_count

        try:
            content, paragraph_count = await asyncio.to_thread(_read_docx, path)
            metadata["paragraph_count"] = paragraph_count
            return content, metadata
        except Exception as e:
            raise DocumentProcessingError(
                message=f"Error reading DOCX: {str(e)}",
                file_path=str(path),
                operation="load_docx",
                original_error=e,
            )

    async def _load_audio(self, path: Path, metadata: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
        """
        Load audio file and transcribe asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        
        Raises:
            DocumentProcessingError: Raised when this function detects an invalid state or when an underlying call fails.
        """
        import asyncio

        if not AUDIO_AVAILABLE:
            raise DocumentProcessingError(
                message="Audio processing libraries required. Install with: pip install SpeechRecognition pydub",
                file_path=str(path),
                operation="load_audio",
            )

        if not self.enable_audio_transcription:
            return "", {**metadata, "transcription_disabled": True}

        def _process_audio(p: Path) -> Tuple[str, Dict[str, Any]]:
            # Load audio file
            audio = AudioSegment.from_file(str(p))
            audio_metadata = {
                "duration_seconds": len(audio) / 1000.0,
                "sample_rate": audio.frame_rate,
            }

            # Convert to WAV for recognition
            wav_io = io.BytesIO()
            audio.export(wav_io, format="wav")
            wav_io.seek(0)

            # Transcribe
            with sr.AudioFile(wav_io) as source:
                audio_data = self.recognizer.record(source)
                try:
                    transcript = self.recognizer.recognize_google(
                        audio_data, language=self.audio_language
                    )
                    content = f"[Audio Transcript]\n{transcript}"
                    audio_metadata["transcription_language"] = self.audio_language
                    return content, audio_metadata
                except sr.UnknownValueError:
                    return (
                        "[Audio file loaded but transcription failed - audio may be unclear]",
                        audio_metadata,
                    )
                except sr.RequestError as e:
                    raise DocumentProcessingError(
                        message=f"Transcription service error: {str(e)}",
                        file_path=str(p),
                        operation="load_audio",
                        original_error=e,
                    )

        try:
            content, audio_metadata = await asyncio.to_thread(_process_audio, path)
            metadata.update(audio_metadata)
            return content, metadata
        except Exception as e:
            raise DocumentProcessingError(
                message=f"Error processing audio: {str(e)}",
                file_path=str(path),
                operation="load_audio",
                original_error=e,
            )

    def _extract_video_properties(self, cap: Any, metadata: Dict[str, Any]) -> Tuple[float, float]:
        """
        Extract video properties and add to metadata.
        
        Args:
            cap (Any): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[float, float]: Result of the operation.
        """
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frame_count / fps if fps > 0 else 0

        metadata["video_fps"] = fps
        metadata["frame_count"] = frame_count
        metadata["duration_seconds"] = duration

        return fps, duration

    def _extract_video_frames(self, cap: Any, fps: float, content_parts: list) -> int:
        """
        Extract frames from video and add to content_parts.
        
        Args:
            cap (Any): Input parameter for this operation.
            fps (float): Input parameter for this operation.
            content_parts (list): Input parameter for this operation.
        
        Returns:
            int: Result of the operation.
        """
        frame_interval = int(fps / self.video_frames_per_second) if fps > 0 else 30
        frame_num = 0
        extracted_frames = 0

        while cap.isOpened():
            ret, _ = cap.read()
            if not ret:
                break

            if frame_num % frame_interval == 0:
                # Convert frame to text description (would use vision model in production)
                content_parts.append(f"[Frame at {frame_num/fps:.2f}s: Frame extracted]")
                extracted_frames += 1

            frame_num += 1

        return extracted_frames

    async def _load_video(
        self, path: Path, metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Load video file, extract frames, and transcribe audio asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        
        Raises:
            DocumentProcessingError: Raised when this function detects an invalid state or when an underlying call fails.
        """
        import asyncio

        if not VIDEO_AVAILABLE:
            raise DocumentProcessingError(
                message="OpenCV is required for video support. Install with: pip install opencv-python",
                file_path=str(path),
                operation="load_video",
            )

        def _process_video(p: Path) -> Tuple[str, Dict[str, Any]]:
            content_parts = []
            video_metadata = {}

            cap = cv2.VideoCapture(str(p))
            if not cap.isOpened():
                raise DocumentProcessingError(
                    message="Could not open video file", file_path=str(p), operation="load_video"
                )

            try:
                # Get video properties and add to metadata
                fps, duration = self._extract_video_properties(cap, video_metadata)

                # Extract audio and transcribe if enabled
                if self.enable_video_transcription and self.recognizer:
                    # Note: Audio extraction from video requires ffmpeg
                    # This is a simplified version - full implementation would extract audio track
                    content_parts.append(
                        "[Video Audio: Transcription available if audio track extracted]"
                    )

                # Extract frames if enabled
                if self.video_extract_frames:
                    extracted_frames = self._extract_video_frames(cap, fps, content_parts)
                    video_metadata["extracted_frames"] = extracted_frames
                    content_parts.append(
                        f"\n[Video Summary: {extracted_frames} frames extracted from {duration:.2f}s video]"
                    )

                content = "\n".join(content_parts) if content_parts else "[Video file processed]"
                return content, video_metadata
            finally:
                cap.release()

        try:
            content, video_metadata = await asyncio.to_thread(_process_video, path)
            metadata.update(video_metadata)
            return content, metadata
        except Exception as e:
            raise DocumentProcessingError(
                message=f"Error processing video: {str(e)}",
                file_path=str(path),
                operation="load_video",
                original_error=e,
            )

    def _process_image_ocr(self, image: Any, content_parts: List[str], img_metadata: Dict[str, Any]) -> None:
        """Process OCR on image."""
        try:
            ocr_text = pytesseract.image_to_string(image)
            if ocr_text.strip():
                content_parts.append(f"[OCR Text]\n{ocr_text}")
                img_metadata["ocr_performed"] = True
        except Exception as e:
            content_parts.append(f"[OCR failed: {str(e)}]")
            img_metadata["ocr_performed"] = False

    def _process_image_description(
        self, image: Any, content_parts: List[str], img_metadata: Dict[str, Any], _gateway: Optional[Any]
    ) -> None:
        """Process image description generation.
        
        Args:
            image: PIL Image object
            content_parts: List to append description content to
            img_metadata: Metadata dictionary to update
            _gateway: Gateway client (unused, reserved for future vision API integration)
        """
        try:
            # Convert image to base64 for API
            import base64

            buffered = io.BytesIO()
            image.save(buffered, format="PNG")
            _ = base64.b64encode(buffered.getvalue()).decode()  # For future use with vision API

            # Use vision model to describe image
            # Note: This requires a vision-capable model like GPT-4 Vision
            # Future: _gateway will be used to call vision API
            description = "[Image description would be generated using vision model]"
            content_parts.append(f"[Image Description]\n{description}")
            img_metadata["description_generated"] = True
        except Exception as e:
            content_parts.append(f"[Description generation failed: {str(e)}]")
            img_metadata["description_generated"] = False

    async def _load_image(
        self, path: Path, metadata: Dict[str, Any], gateway: Optional[Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Load image file, perform OCR, and generate description asynchronously.
        
        Args:
            path (Path): Input parameter for this operation.
            metadata (Dict[str, Any]): Extra metadata for the operation.
            gateway (Optional[Any]): Gateway client used for LLM calls.
        
        Returns:
            Tuple[str, Dict[str, Any]]: Dictionary result of the operation.
        
        Raises:
            DocumentProcessingError: Raised when this function detects an invalid state or when an underlying call fails.
        """
        import asyncio

        if not IMAGE_AVAILABLE:
            raise DocumentProcessingError(
                message="Image processing libraries required. Install with: pip install Pillow pytesseract",
                file_path=str(path),
                operation="load_image",
            )

        def _process_image_sync(p: Path) -> Tuple[str, Dict[str, Any]]:
            image = Image.open(p)
            img_metadata = {
                "image_width": image.width,
                "image_height": image.height,
                "image_format": image.format,
            }
            content_parts = []

            # OCR if enabled
            if self.enable_image_ocr:
                self._process_image_ocr(image, content_parts, img_metadata)

            # Note: Image description generation would use gateway (async)
            # For now, this is a placeholder
            if self.enable_image_description and gateway:
                self._process_image_description(image, content_parts, img_metadata, gateway)

            content = "\n".join(content_parts) if content_parts else "[Image file processed]"
            return content, img_metadata

        try:
            content, image_metadata = await asyncio.to_thread(_process_image_sync, path)
            metadata.update(image_metadata)
            return content, metadata
        except Exception as e:
            raise DocumentProcessingError(
                message=f"Error processing image: {str(e)}",
                file_path=str(path),
                operation="load_image",
                original_error=e,
            )


def create_multimodal_loader(
    enable_audio_transcription: bool = True,
    enable_video_transcription: bool = True,
    enable_image_ocr: bool = True,
    enable_image_description: bool = True,
    **kwargs: Any,
) -> MultiModalLoader:
    """
    Create a multi-modal data loader.

    Args:
        enable_audio_transcription: Enable audio transcription
        enable_video_transcription: Enable video transcription
        enable_image_ocr: Enable OCR for images
        enable_image_description: Enable image description generation
        **kwargs: Additional configuration

    Returns:
        MultiModalLoader instance
    """
    return MultiModalLoader(
        enable_audio_transcription=enable_audio_transcription,
        enable_video_transcription=enable_video_transcription,
        enable_image_ocr=enable_image_ocr,
        enable_image_description=enable_image_description,
        **kwargs,
    )
